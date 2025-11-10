import os
import dotenv
import traceback
import requests
from docx import Document
from PyQt5.QtWidgets import QApplication, QMainWindow, QFileDialog
from PyQt5.QtCore import QObject, QThread, pyqtSignal
from PyQt5.QtGui import QPixmap
from main_window_ui import Ui_MainWindow
from openai import OpenAI
from filesearch.file import run_file_search  # ✅ file_search 연결

dotenv.load_dotenv()


# --- Rudebot Worker ---
class RudebotWorker(QObject):
    finished = pyqtSignal(str)
    error = pyqtSignal(str)

    def __init__(self, client, question):
        super().__init__()
        self.client = client
        self.question = question

    def run(self):
        try:
            response = self.client.chat.completions.create(
                model="ft:gpt-3.5-turbo-0125:personal::CaKAw4RI",  # ✅ 튜닝된 모델 ID
                messages=[
                    {"role": "system", "content": "You are RudeBot — a sarcastic chatbot."},
                    {"role": "user", "content": self.question}
                ]
            )
            text = response.choices[0].message.content
            self.finished.emit(text)
        except Exception as e:
            traceback.print_exc()
            self.error.emit(str(e))


# --- Image Worker ---
class ImageWorker(QObject):
    finished = pyqtSignal(str)
    error = pyqtSignal(str)

    def __init__(self, client, prompt):
        super().__init__()
        self.client = client
        self.prompt = prompt

    def run(self):
        try:
            response = self.client.images.generate(
                model="dall-e-3",
                prompt=self.prompt,
                n=1,
                size="1024x1024",
                response_format="url"
            )
            image_url = response.data[0].url
            self.finished.emit(image_url)
        except Exception as e:
            traceback.print_exc()
            self.error.emit(str(e))


# --- Audio Worker ---
class AudioWorker(QObject):
    finished = pyqtSignal(dict, str)
    error = pyqtSignal(str)

    def __init__(self, client, audio_file_path, output_filename):
        super().__init__()
        self.client = client
        self.audio_file_path = audio_file_path
        self.output_filename = output_filename

    def run(self):
        try:
            with open(self.audio_file_path, "rb") as audio_file:
                transcript = self.client.audio.transcriptions.create(
                    file=audio_file,
                    model="whisper-1",
                    response_format="text"
                )

            transcription_text = transcript
            notes = {
                "abstract_summary": self.abstract_summary_extraction(transcription_text),
                "key_points": self.key_points_extraction(transcription_text),
                "action_items": self.action_items_extraction(transcription_text),
                "sentiment": self.sentiment_analysis(transcription_text)
            }

            doc = Document()
            for key, value in notes.items():
                heading = ''.join(word.capitalize() for word in key.split('_'))
                doc.add_heading(heading, level=1)
                doc.add_paragraph(value)
                doc.add_paragraph()
            doc.save(self.output_filename)

            self.finished.emit(notes, self.output_filename)

        except Exception as e:
            traceback.print_exc()
            self.error.emit(str(e))

    def abstract_summary_extraction(self, text):
        return self._ask_gpt("Summarize the following text:", text)

    def key_points_extraction(self, text):
        return self._ask_gpt("Extract key points:", text)

    def action_items_extraction(self, text):
        return self._ask_gpt("Extract action items:", text)

    def sentiment_analysis(self, text):
        return self._ask_gpt("Analyze the sentiment:", text)

    def _ask_gpt(self, instruction, text):
        response = self.client.chat.completions.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": "You are a helpful assistant."},
                {"role": "user", "content": f"{instruction}\n\n{text}"}
            ]
        )
        return response.choices[0].message.content


# --- Main Window ---
class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.ui = Ui_MainWindow()
        self.ui.setupUi(self)
        self.client = OpenAI(api_key=os.getenv("API_KEY"))

        # 메뉴 전환
        self.ui.menu_list.currentRowChanged.connect(self.ui.stackedWidget.setCurrentIndex)

        # 기능 버튼 연결
        self.ui.poem_generate_btn.clicked.connect(self.generate_poem)
        self.ui.translate_btn.clicked.connect(self.translate_text)
        self.ui.image_generate_btn.clicked.connect(self.generate_image)
        self.ui.audio_note_btn.clicked.connect(self.generate_audio_note)
        self.ui.file_btn.clicked.connect(self.handle_file_search)
        self.ui.rudebot_btn_2.clicked.connect(self.generate_rudebot)  # ✅ 루드봇 버튼

        self.show()

    # --- 시 생성 ---
    def generate_poem(self):
        topic = self.ui.poem_topic_input.text()
        if not topic:
            self.ui.poem_result_view.setText("시의 주제를 입력해주세요.")
            return

        self.ui.poem_generate_btn.setEnabled(False)
        self.ui.poem_result_view.setText("시를 생성 중입니다...")

        try:
            completion = self.client.chat.completions.create(
                model="gpt-3.5-turbo",
                messages=[
                    {"role": "system", "content": "You are a helpful assistant who writes poems in Korean."},
                    {"role": "user", "content": f"{topic}라는 주제로 시를 써줘."}
                ]
            )
            result_text = completion.choices[0].message.content
            self.ui.poem_result_view.setText(result_text)
        except Exception as e:
            self.ui.poem_result_view.setText(f"오류 발생: {e}")
        finally:
            self.ui.poem_generate_btn.setEnabled(True)

    # --- 번역 ---
    def translate_text(self):
        source_text = self.ui.translate_source_input.toPlainText()
        if not source_text:
            self.ui.translate_result_view.setText("번역할 문장을 입력해주세요.")
            return

        self.ui.translate_result_view.setText("번역 중입니다...")
        try:
            completion = self.client.chat.completions.create(
                model="gpt-3.5-turbo",
                messages=[
                    {"role": "system", "content": "Translate English to Korean."},
                    {"role": "user", "content": source_text}
                ]
            )
            result = completion.choices[0].message.content
            self.ui.translate_result_view.setText(result)
        except Exception as e:
            self.ui.translate_result_view.setText(f"오류 발생: {e}")

    # --- 이미지 생성 ---
    def generate_image(self):
        prompt = self.ui.image_prompt_input.text()
        if not prompt:
            self.ui.image_display_label.setText("이미지 설명을 입력해주세요.")
            return

        self.ui.image_display_label.setText("이미지 생성 중...")
        self.image_thread = QThread()
        self.image_worker = ImageWorker(self.client, prompt)
        self.image_worker.moveToThread(self.image_thread)

        self.image_thread.started.connect(self.image_worker.run)
        self.image_worker.finished.connect(self.handle_image_result)
        self.image_worker.error.connect(self.handle_image_error)
        self.image_worker.finished.connect(self.image_thread.quit)
        self.image_thread.start()

    def handle_image_result(self, url):
        try:
            data = requests.get(url).content
            pixmap = QPixmap()
            pixmap.loadFromData(data)
            self.ui.image_display_label.setPixmap(pixmap)
        except Exception as e:
            self.ui.image_display_label.setText(f"로드 실패: {e}")

    def handle_image_error(self, msg):
        self.ui.image_display_label.setText(f"오류: {msg}")

    # --- 오디오 노트 생성 ---
    def generate_audio_note(self):
        path = self.ui.audio_source_input.text()
        if not path:
            self.ui.audio_source_input.setText("오디오 파일 경로를 입력하세요.")
            return

        output_file = QFileDialog.getSaveFileName(self, "저장할 파일", "", "Word Document (*.docx)")[0]
        if not output_file:
            return

        self.ui.audio_source_input.setText("노트 생성 중...")
        self.audio_thread = QThread()
        self.audio_worker = AudioWorker(self.client, path, output_file)
        self.audio_worker.moveToThread(self.audio_thread)

        self.audio_thread.started.connect(self.audio_worker.run)
        self.audio_worker.finished.connect(self.handle_audio_result)
        self.audio_worker.error.connect(self.handle_audio_error)
        self.audio_thread.start()

    def handle_audio_result(self, notes, filename):
        self.ui.audio_source_input.setText(f"노트 생성 완료: {filename}")

    def handle_audio_error(self, err):
        self.ui.audio_source_input.setText(f"오류 발생: {err}")

    # --- 파일 검색 ---
    def handle_file_search(self):
        file_path = self.ui.file_input.text()
        question = self.ui.user_input.text()
        if not file_path or not question:
            self.ui.translate_result_view_2.setText("파일 경로와 질문을 입력해주세요.")
            return

        self.ui.translate_result_view_2.setText("파일 분석 중입니다...")

        try:
            result = run_file_search(file_path, question, self.client)
            self.ui.translate_result_view_2.setText(result)
        except Exception as e:
            self.ui.translate_result_view_2.setText(f"오류 발생: {e}")

    # --- 루드봇 ---
    def generate_rudebot(self):
        question = self.ui.rudebot_input_2.text().strip()
        if not question:
            self.ui.translate_result_view_3.setText("질문을 입력해주세요.")
            return

        self.ui.translate_result_view_3.setText("루드봇이 생각 중입니다... 🤔")

        self.rudebot_thread = QThread()
        self.rudebot_worker = RudebotWorker(self.client, question)
        self.rudebot_worker.moveToThread(self.rudebot_thread)

        self.rudebot_thread.started.connect(self.rudebot_worker.run)
        self.rudebot_worker.finished.connect(self.handle_rudebot_result)
        self.rudebot_worker.error.connect(self.handle_rudebot_error)
        self.rudebot_worker.finished.connect(self.rudebot_thread.quit)
        self.rudebot_thread.start()

    def handle_rudebot_result(self, text):
        self.ui.translate_result_view_3.setText(f"{text}\n\n✅ 실행 완료 (루드봇 응답)")

    def handle_rudebot_error(self, msg):
        self.ui.translate_result_view_3.setText(f"⚠️ 오류 발생: {msg}")


if __name__ == "__main__":
    import sys
    app = QApplication(sys.argv)
    window = MainWindow()
    sys.exit(app.exec())
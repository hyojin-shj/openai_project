import traceback
from docx import Document
from PyQt5.QtCore import QObject, pyqtSignal
import os


class AudioWorker(QObject):
    finished = pyqtSignal(dict, str)
    error = pyqtSignal(str)

    def __init__(self, client, audio_file_path, output_filename):
        super().__init__()
        self.client = client
        self.audio_file_path = audio_file_path
        self.output_filename = output_filename

    def run(self):
        try:
            with open(self.audio_file_path, "rb") as audio_file:
                transcript = self.client.audio.transcriptions.create(
                    file=audio_file,
                    model="whisper-1",
                    response_format="text"
                )

            if hasattr(transcript, "text"):
                transcription_text = transcript.text
            else:
                transcription_text = str(transcript)

            output_dir = os.path.dirname(self.output_filename)
            if output_dir and not os.path.exists(output_dir):
                os.makedirs(output_dir, exist_ok=True)

            notes = {
                "abstract_summary": self.abstract_summary_extraction(transcription_text),
                "key_points": self.key_points_extraction(transcription_text),
                "action_items": self.action_items_extraction(transcription_text),
                "sentiment": self.sentiment_analysis(transcription_text)
            }

            doc = Document()
            for key, value in notes.items():
                heading = ''.join(word.capitalize() for word in key.split('_'))
                doc.add_heading(heading, level=1)
                doc.add_paragraph(value)
                doc.add_paragraph()
            doc.save(self.output_filename)

            self.finished.emit(notes, self.output_filename)

        except Exception as e:
            traceback.print_exc()
            self.error.emit(str(e))

    def abstract_summary_extraction(self, text):
        return self._ask_gpt("Summarize the following text:", text)

    def key_points_extraction(self, text):
        return self._ask_gpt("Extract key points:", text)

    def action_items_extraction(self, text):
        return self._ask_gpt("Extract action items:", text)

    def sentiment_analysis(self, text):
        return self._ask_gpt("Analyze the sentiment:", text)

    def _ask_gpt(self, instruction, text):
        response = self.client.chat.completions.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": "You are a helpful assistant."},
                {"role": "user", "content": f"{instruction}\n\n{text}"}
            ]
        )
        return response.choices[0].message.content